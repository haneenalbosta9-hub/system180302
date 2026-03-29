import os
import re
import json
from datetime import date, datetime, timedelta
from typing import Optional

import pandas as pd
import streamlit as st
from docx import Document
import gspread
from google.oauth2.service_account import Credentials

# =====================================================
# Configuration
# =====================================================
REPORTS_DIR = "reports"
BACKUP_DIR = "file_backups"
os.makedirs(BACKUP_DIR, exist_ok=True)
TEMPLATE_BIOBURDEN = "BioburdenReport1803.docx"
TEMPLATE_STERILITY = "SterilityReport.docx"
TEMPLATE_ENDOTOXIN = "EndotoxinReport.docx"
TEMPLATE_ENVIRONMENT = "EnvironmentReport.docx"

# =====================================================
# Media Integration — links to microbiologymicrobiology_store.py files
# =====================================================
MEDIA_PREP_FILE = "media_preparation.xlsx"
MEDIA_INVENTORY_FILE = "microbiology_inventory.xlsx"

# Mapping: Test Type → list of media keywords to search for in media_preparation
# Keywords are matched case-insensitively against "Media Type" column
TEST_MEDIA_MAP = {
    "Bioburden": [
        {"label": "TAMC (Aerobic Count)", "keywords": ["tryptone soya agar", "tsa", "casein soya", "scda", "nutrient agar"]},
        {"label": "TYMC (Yeast/Mould Count)", "keywords": ["sabouraud", "sda", "sab dex", "malt extract agar"]},
    ],
    "Sterility": [
        {"label": "Aerobic / Facultative Anaerobic (FTM)", "keywords": ["thioglycollate", "ftm", "fluid thioglycolate"]},
        {"label": "Aerobic Organisms (SCDM)", "keywords": ["tryptone soya broth", "tsb", "casein soya broth", "scdm", "scd broth"]},
    ],
    "Environmental": [
        {"label": "Settle Plates / Contact Plates (TSA)", "keywords": ["tryptone soya agar", "tsa", "casein soya", "scda"]},
        {"label": "Yeast/Mould Plates (SDA)", "keywords": ["sabouraud", "sda", "malt extract agar"]},
    ],
    "Endotoxin": [],  # No culture media — LAL-based test
}

# Default mL consumed per test per media type (can be overridden in UI)
DEFAULT_MEDIA_CONSUMPTION_ML = {
    "Bioburden": 15,
    "Sterility": 100,
    "Environmental": 20,
    "Endotoxin": 0,
}


@st.cache_data(ttl=30)
def load_media_prep_cached():
    """Load media_preparation.xlsx with a Volume Consumed column."""
    if os.path.exists(MEDIA_PREP_FILE):
        df = pd.read_excel(MEDIA_PREP_FILE)
        if "Volume Consumed (mL)" not in df.columns:
            df["Volume Consumed (mL)"] = 0.0
        df["Volume Consumed (mL)"] = pd.to_numeric(
            df["Volume Consumed (mL)"], errors="coerce"
        ).fillna(0.0)
        df["Quantity (mL)"] = pd.to_numeric(
            df["Quantity (mL)"], errors="coerce"
        ).fillna(0.0)
        df["Volume Remaining (mL)"] = df["Quantity (mL)"] - df["Volume Consumed (mL)"]
        if "Date" in df.columns:
            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        if "Expiry Date" in df.columns:
            df["Expiry Date"] = pd.to_datetime(df["Expiry Date"], errors="coerce")
        return df
    return pd.DataFrame(columns=[
        "Date", "Media Type", "Lot No.", "Quantity (mL)", "Media Used (g)",
        "Water Used (mL)", "Prepared By", "Expiry Date", "Sterilization Method",
        "Notes", "Volume Consumed (mL)", "Volume Remaining (mL)"
    ])


def save_media_prep_with_consumption(df):
    """Persist media preparation records (including Volume Consumed column)."""
    save_df = df.drop(columns=["Volume Remaining (mL)"], errors="ignore")
    save_df.to_excel(MEDIA_PREP_FILE, index=False)
    # Bust cache
    load_media_prep_cached.clear()


def get_suggested_media_for_test(test_type: str):
    """
    For a given test type return a list of dicts, each describing one required media:
      {label, keywords, latest_batch_row, remaining_ml, has_stock}
    """
    media_reqs = TEST_MEDIA_MAP.get(test_type, [])
    if not media_reqs:
        return []

    df_prep = load_media_prep_cached()
    today = pd.Timestamp(date.today())
    results = []

    for req in media_reqs:
        label = req["label"]
        keywords = req["keywords"]

        # Filter prep records where Media Type matches any keyword
        if not df_prep.empty:
            mask = df_prep["Media Type"].str.lower().apply(
                lambda x: any(kw in str(x).lower() for kw in keywords)
            )
            matching = df_prep[mask].copy()
        else:
            matching = pd.DataFrame()

        # Exclude expired batches (Expiry Date in the past)
        if not matching.empty and "Expiry Date" in matching.columns:
            matching = matching[
                matching["Expiry Date"].isna() | (matching["Expiry Date"] >= today)
            ]

        # Keep only batches with remaining volume > 0
        if not matching.empty:
            with_stock = matching[matching["Volume Remaining (mL)"] > 0]
        else:
            with_stock = pd.DataFrame()

        if not with_stock.empty:
            # Latest prepared batch = highest Date
            latest = with_stock.sort_values("Date", ascending=False).iloc[0]
            results.append({
                "label": label,
                "keywords": keywords,
                "latest_batch": latest,
                "remaining_ml": float(latest["Volume Remaining (mL)"]),
                "has_stock": True,
            })
        else:
            # No batch with remaining stock
            results.append({
                "label": label,
                "keywords": keywords,
                "latest_batch": None,
                "remaining_ml": 0.0,
                "has_stock": False,
            })

    return results


def deduct_media_consumption(media_type_keywords: list, volume_ml: float):
    """
    Deduct volume_ml from the latest non-expired batch that matches any keyword.
    Returns (success: bool, message: str).
    """
    df_prep = load_media_prep_cached()
    if df_prep.empty:
        return False, "No media preparation records found."

    today = pd.Timestamp(date.today())
    mask = df_prep["Media Type"].str.lower().apply(
        lambda x: any(kw in str(x).lower() for kw in media_type_keywords)
    )
    matching_idx = df_prep[mask & (df_prep["Volume Remaining (mL)"] > 0)].index
    if not df_prep.empty and "Expiry Date" in df_prep.columns:
        valid_mask = mask & (
            (df_prep["Volume Remaining (mL)"] > 0) &
            (df_prep["Expiry Date"].isna() | (df_prep["Expiry Date"] >= today))
        )
        matching_idx = df_prep[valid_mask].index

    if matching_idx.empty:
        return False, "No batch with remaining stock found."

    # Pick the latest by Date
    latest_idx = df_prep.loc[matching_idx].sort_values("Date", ascending=False).index[0]
    old_consumed = float(df_prep.at[latest_idx, "Volume Consumed (mL)"])
    df_prep.at[latest_idx, "Volume Consumed (mL)"] = old_consumed + volume_ml
    save_media_prep_with_consumption(df_prep)
    return True, f"Deducted {volume_ml} mL from batch Lot {df_prep.at[latest_idx, 'Lot No.']}."


# ✅ Your Google Sheet ID hardcoded — no spreadsheet_name in secrets needed
SPREADSHEET_ID = "1EXiXsOQ0VsfIbZlUpN3r6g0-aRNUUEKZDVZHh_xZnEY"

COLUMNS = [
    "Received Date", "Sample ID", "Unit No.", "Sample Type", "Sample Batch No.",
    "Customer Name", "Reference No.", "Type of Test",
    "Test Performing Date", "Test Status", "Product Name",
    "Customer Name (AR)", "Customer Name (EN)"
]

TEST_TYPES = [
    "Bioburden", "Endotoxin", "Sterility", "Environmental",
    "Total Coliforms & E. Coli", "Pseudomonas aeruginosa",
    "Total heterotrophic bacterial count", "Legionella", "Fungi",
    "Other (Not Listed)"
]

SAMPLE_TYPES_DEFAULT = [
    "SAFEPIT 135 µm", "SAFEPIT 145 µm", "SAFEPIT 170 µm", "SAFEPIT 300 µm",
    "SAFEPIT 200 µm", "SAFEPIT 80 µm", "SAFEPIT 155 µm", "SAFEPIT 290 µm",
    "Pipette 140 µm", "Pipette 170 µm", "Pipette 300 µm", "Capillary tube",
    "Settle Plates", "AccurET Catheter", "Single Lumen",
    "Drinking Water", "Swimming Water", "Dialysis Water", "Other (Not listed)"
]

SHEET_SAMPLES = "Samples"
SHEET_SAMPLE_TYPES = "SampleTypes"
SHEET_TEST_TYPES = "TestTypes"
SHEET_CUSTOMERS_EN = "CustomersEN"
SHEET_CUSTOMERS_AR = "CustomersAR"

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

# =====================================================
# Google Sheets — opens by SPREADSHEET_ID, not by name
# =====================================================


def _get_worksheet(sheet_name, headers=None):
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    gc = gspread.authorize(creds)
    ss = gc.open_by_key(SPREADSHEET_ID)
    try:
        ws = ss.worksheet(sheet_name)
    except gspread.WorksheetNotFound:
        ws = ss.add_worksheet(title=sheet_name, rows=2000, cols=30)
        if headers:
            ws.append_row(headers)
    return ws


def _df_to_rows(df):
    df_save = df.copy()
    for col in df_save.columns:
        if pd.api.types.is_datetime64_any_dtype(df_save[col]):
            # datetime64 column — format directly
            df_save[col] = df_save[col].dt.strftime("%Y-%m-%d")
        else:
            # object column — may contain Python date/datetime objects
            def _fmt(v):
                if hasattr(v, 'strftime'):
                    return v.strftime("%Y-%m-%d")
                if v is None or (isinstance(v, float) and pd.isna(v)):
                    return ""
                return str(v)
            df_save[col] = df_save[col].map(_fmt)
    cols_to_write = [c for c in COLUMNS if c in df_save.columns]
    return df_save[cols_to_write].values.tolist()

# =====================================================
# Data read/write
# =====================================================


@st.cache_data(ttl=60)
def get_data():
    try:
        ws = _get_worksheet(SHEET_SAMPLES, headers=COLUMNS)
        records = ws.get_all_records(value_render_option="FORMATTED_VALUE")
        if not records:
            return pd.DataFrame(columns=COLUMNS)
        df = pd.DataFrame(records)
        for col in COLUMNS:
            if col not in df.columns:
                df[col] = ""

        def _parse_date_col(series):
            """Parse dates that may be YYYY-MM-DD, DD/MM/YYYY, or M/D/YYYY (Sheets display)."""
            # Try ISO format first (most reliable), then fall back
            parsed = pd.to_datetime(
                series, errors="coerce", dayfirst=False, format="%Y-%m-%d")
            # For values that failed ISO, try dayfirst formats
            mask = parsed.isna() & series.astype(str).str.strip().ne("")
            if mask.any():
                parsed[mask] = pd.to_datetime(
                    series[mask], errors="coerce", dayfirst=True)
            return parsed

        df["Received Date"] = _parse_date_col(df["Received Date"])
        df["Test Performing Date"] = _parse_date_col(
            df["Test Performing Date"])
        df["Unit No."] = pd.to_numeric(
            df["Unit No."], errors="coerce").fillna(1).astype(int)
        return df
    except Exception as e:
        st.error(f"❌ Error reading from Google Sheets: {e}")
        return pd.DataFrame(columns=COLUMNS)


def save_data(df):
    """
    FULL REWRITE — only call this for delete operations or first-time setup.
    For adds use append_rows(); for edits use update_rows_targeted().
    """
    try:
        get_data.clear()
        load_custom_lists.clear()
        ws = _get_worksheet(SHEET_SAMPLES, headers=COLUMNS)

        # Safety check before clearing
        existing_headers = ws.row_values(1)
        if existing_headers and existing_headers != COLUMNS:
            st.error(
                "❌ SAFETY CHECK FAILED: Sheet headers don't match expected columns!")
            st.error(f"Expected: {COLUMNS}")
            st.error(f"Found: {existing_headers}")
            st.warning(
                "⚠️ Please fix the column headers in your Google Sheet before saving.")
            st.stop()

        rows = _df_to_rows(df)
        ws.clear()
        ws.append_row(COLUMNS)
        if rows:
            ws.append_rows(rows, value_input_option="RAW")
        st.toast("✅ Saved to Google Sheets", icon="✅")
    except Exception as e:
        st.error(f"❌ Error saving to Google Sheets: {e}")


def append_rows(new_rows_df):
    """
    SAFE ADD — only appends new rows, never touches existing data.
    Use this for all new-sample operations.
    """
    try:
        get_data.clear()
        ws = _get_worksheet(SHEET_SAMPLES, headers=COLUMNS)
        rows = _df_to_rows(new_rows_df)
        if rows:
            ws.append_rows(rows, value_input_option="RAW")
        st.toast("✅ Saved to Google Sheets", icon="✅")
    except Exception as e:
        st.error(f"❌ Error appending to Google Sheets: {e}")


def update_rows_targeted(changed_df):
    """
    SAFE EDIT — finds existing rows by Sample ID + Unit No. + Sample Batch No.
    and updates only those specific cells. Never clears the sheet.
    Falls back to save_data only if a row cannot be found (it is new).
    """
    try:
        get_data.clear()
        ws = _get_worksheet(SHEET_SAMPLES, headers=COLUMNS)
        all_values = ws.get_all_values()

        if len(all_values) < 2:
            # Sheet is empty — just save everything
            save_data(changed_df)
            return

        headers = all_values[0]
        # Build column-index lookup for the key fields

        def col_idx(name):
            return headers.index(name) if name in headers else -1

        sid_i = col_idx("Sample ID")
        unit_i = col_idx("Unit No.")
        batch_i = col_idx("Sample Batch No.")

        def row_key(sheet_row):
            return (
                str(sheet_row[sid_i]).strip() if sid_i >= 0 else "",
                str(sheet_row[unit_i]).strip() if unit_i >= 0 else "",
                str(sheet_row[batch_i]).strip() if batch_i >= 0 else "",
            )

        # Map each key → 1-based sheet row number (offset +2 because row 1 is header)
        key_to_sheet_row = {}
        for i, sheet_row in enumerate(all_values[1:], start=2):
            key_to_sheet_row[row_key(sheet_row)] = i

        batch_updates = []
        rows_not_found = []

        for _, row_data in changed_df.iterrows():
            key = (
                str(row_data.get("Sample ID",        "")).strip(),
                str(row_data.get("Unit No.",          "")).strip(),
                str(row_data.get("Sample Batch No.", "")).strip(),
            )
            sheet_row_num = key_to_sheet_row.get(key)

            # Build the cell values for this row
            row_values = []
            for col in COLUMNS:
                val = row_data.get(col, "")
                try:
                    if hasattr(val, "strftime"):
                        val = val.strftime("%Y-%m-%d") if pd.notna(val) else ""
                    elif pd.isna(val):
                        val = ""
                except (TypeError, ValueError):
                    pass
                row_values.append(str(val))

            if sheet_row_num is not None:
                end_col = chr(64 + len(COLUMNS))  # e.g. 'M' for 13 columns
                batch_updates.append({
                    "range": f"A{sheet_row_num}:{end_col}{sheet_row_num}",
                    "values": [row_values],
                })
            else:
                rows_not_found.append(row_values)

        if batch_updates:
            ws.batch_update(batch_updates, value_input_option="RAW")

        # Any rows not found in the sheet are truly new — append them
        if rows_not_found:
            ws.append_rows(rows_not_found, value_input_option="RAW")

        st.toast("✅ Updated Google Sheets", icon="✅")
    except Exception as e:
        st.error(f"❌ Error updating Google Sheets: {e}")


def delete_rows_targeted(sample_id, unit_no=None, batch_no=None):
    """
    SAFE DELETE — finds the exact rows in the sheet by Sample ID (and optionally
    Unit No. / Batch No.) and deletes only those rows.
    Never clears the rest of the sheet.
    """
    try:
        get_data.clear()
        ws = _get_worksheet(SHEET_SAMPLES, headers=COLUMNS)
        all_values = ws.get_all_values()

        if len(all_values) < 2:
            return 0  # Nothing to delete

        headers = all_values[0]

        def col_idx(name):
            return headers.index(name) if name in headers else -1

        sid_i = col_idx("Sample ID")
        unit_i = col_idx("Unit No.")

        # Collect 1-based row numbers to delete (in reverse order so indices stay valid)
        rows_to_delete = []
        for i, row in enumerate(all_values[1:], start=2):
            if sid_i < 0:
                continue
            if str(row[sid_i]).strip() != str(sample_id).strip():
                continue
            if unit_no is not None and unit_i >= 0:
                if str(row[unit_i]).strip() != str(unit_no).strip():
                    continue
            rows_to_delete.append(i)

        if not rows_to_delete:
            return 0

        # Delete from bottom to top so row numbers stay accurate
        for row_num in sorted(rows_to_delete, reverse=True):
            ws.delete_rows(row_num)

        st.toast(f"✅ Deleted {len(rows_to_delete)} row(s)", icon="✅")
        return len(rows_to_delete)
    except Exception as e:
        st.error(f"❌ Error deleting from Google Sheets: {e}")
        return 0


# =====================================================
# Custom Lists
# =====================================================


def _read_list_sheet(sheet_name):
    try:
        ws = _get_worksheet(sheet_name, headers=["value"])
        records = ws.get_all_records()
        return [str(r["value"]) for r in records if r.get("value")]
    except:
        return []


def _write_list_sheet(sheet_name, values):
    try:
        ws = _get_worksheet(sheet_name, headers=["value"])
        ws.clear()
        ws.append_row(["value"])
        if values:
            ws.append_rows([[v] for v in values], value_input_option="RAW")
    except Exception as e:
        st.error(f"❌ Error writing {sheet_name}: {e}")


@st.cache_data(ttl=300)
def load_custom_lists():
    return (
        _read_list_sheet(SHEET_SAMPLE_TYPES),
        _read_list_sheet(SHEET_TEST_TYPES),
        _read_list_sheet(SHEET_CUSTOMERS_EN),
        _read_list_sheet(SHEET_CUSTOMERS_AR),
    )


def save_custom_lists(sample_types, test_types, customers_en, customers_ar):
    _write_list_sheet(SHEET_SAMPLE_TYPES, sample_types)
    _write_list_sheet(SHEET_TEST_TYPES,   test_types)
    _write_list_sheet(SHEET_CUSTOMERS_EN, customers_en)
    _write_list_sheet(SHEET_CUSTOMERS_AR, customers_ar)


def add_custom_value(list_name, new_value):
    sample_types, test_types, customers_en, customers_ar = load_custom_lists()
    changed = False
    if list_name == "SampleTypes" and new_value not in sample_types:
        sample_types.append(new_value)
        changed = True
    if list_name == "TestTypes" and new_value not in test_types:
        test_types.append(new_value)
        changed = True
    if list_name == "CustomersEN" and new_value not in customers_en:
        customers_en.append(new_value)
        changed = True
    if list_name == "CustomersAR" and new_value not in customers_ar:
        customers_ar.append(new_value)
        changed = True
    if changed:
        save_custom_lists(sample_types, test_types, customers_en, customers_ar)
        load_custom_lists.clear()


def get_sample_types():
    custom, _, _, _ = load_custom_lists()
    base = [t for t in SAMPLE_TYPES_DEFAULT if t != "Other (Not listed)"]
    extras = [t for t in custom if t not in base]
    return base + extras + ["Other (Not listed)"]


def get_test_types():
    _, custom, _, _ = load_custom_lists()
    base = [t for t in TEST_TYPES if t != "Other (Not Listed)"]
    extras = [t for t in custom if t not in base]
    return base + extras + ["Other (Not Listed)"]


def get_customers_en():
    _, _, custom, _ = load_custom_lists()
    return sorted(set(custom))


def get_customers_ar():
    _, _, _, custom = load_custom_lists()
    return sorted(set(custom))

# =====================================================
# Helpers
# =====================================================


def get_next_unit_no(df, sample_id, batch_no):
    mask = (df["Sample ID"] == sample_id) & (
        df["Sample Batch No."].astype(str) == str(batch_no))
    existing = df.loc[mask, "Unit No."]
    return 1 if existing.empty else int(existing.max()) + 1


def get_next_serial(df, batch_queue=None):
    serials = []
    if not df.empty:
        for sid in df["Sample ID"].dropna():
            m = re.match(r"MIC-(\d+)-", str(sid))
            if m:
                serials.append(int(m.group(1)))
    if batch_queue:
        for entry in batch_queue:
            m = re.match(r"MIC-(\d+)-", str(entry.get("Sample ID", "")))
            if m:
                serials.append(int(m.group(1)))
    return (max(serials) + 1) if serials else 1


def format_report_date(date_val):
    if pd.isna(date_val):
        return ""
    if isinstance(date_val, str):
        date_val = pd.to_datetime(date_val)
    return date_val.strftime("%d/%m/%Y")


def safe_report_filename(sample_id):
    return sample_id.replace("/", "_")


def report_path(filename):
    os.makedirs(REPORTS_DIR, exist_ok=True)
    return os.path.join(REPORTS_DIR, filename)


def find_report_path(filename):
    path = report_path(filename)
    return path if os.path.exists(path) else None


def convert_docx_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        import subprocess
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                        "--outdir", os.path.dirname(docx_path), docx_path],
                       capture_output=True, timeout=30)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass
    return None


def replace_placeholders_in_tables(doc, placeholders):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in placeholders.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


def generate_sample_id_range(start_id, end_id):
    try:
        sp = start_id.split("-")
        ep = end_id.split("-")
        return [f"MIC-{i:04d}-{sp[2]}-{sp[3]}" for i in range(int(sp[1]), int(ep[1]) + 1)]
    except:
        return []


# =====================================================
# Page config
# =====================================================
st.set_page_config(page_title="system1803", page_icon="🧪",
                   layout="wide", initial_sidebar_state="expanded")

try:
    SAMPLE_TYPES = get_sample_types()
    TEST_TYPES = get_test_types()
except Exception as _e:
    st.error(f"⚠️ Cannot connect to Google Sheets: {_e}")
    SAMPLE_TYPES = SAMPLE_TYPES_DEFAULT
    TEST_TYPES = ["Bioburden", "Endotoxin", "Sterility", "Environmental",
                  "Total Coliforms & E. Coli", "Pseudomonas aeruginosa",
                  "Total heterotrophic bacterial count", "Legionella", "Fungi", "Other (Not Listed)"]

st.image("header.png", use_container_width=True)
st.markdown("---")
st.sidebar.title("🧪 system1803")
st.sidebar.markdown("---")
menu = st.sidebar.radio("Navigation",
                        ["Dashboard", "Add Sample", "Edit Sample",
                            "Perform Test", "Enter Results"],
                        label_visibility="collapsed")
st.sidebar.markdown("---")

if menu == "Dashboard":
    st.title("Dashboard")
elif menu == "Edit Sample":
    st.title("✏️ Edit Sample")
else:
    st.title("system1803")

# =====================================================
# DASHBOARD
# =====================================================
if menu == "Dashboard":
    df = get_data()
    if df.empty:
        st.info("No data available yet. Start by adding samples!")
    else:
        st.subheader("📊 Dashboard Overview")
        col1, col2 = st.columns(2)
        with col1:
            start_date_1 = st.date_input(
                "Start Date (Chart 1)", value=date(date.today().year, 1, 1))
        with col2:
            end_date_1 = st.date_input(
                "End Date (Chart 1)",   value=date.today())
        col1, col2 = st.columns(2)
        with col1:
            start_date_2 = st.date_input(
                "Start Date (Chart 2)", value=date(date.today().year, 1, 1))
        with col2:
            end_date_2 = st.date_input(
                "End Date (Chart 2)",   value=date.today())

        df_filtered_1 = df[(df["Received Date"] >= pd.Timestamp(start_date_1)) &
                           (df["Received Date"] <= pd.Timestamp(end_date_1) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1))]
        df_filtered_2 = df[(df["Received Date"] >= pd.Timestamp(start_date_2)) &
                           (df["Received Date"] <= pd.Timestamp(end_date_2) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1))]

        st.markdown("---")
        st.subheader("Chart 1: Test Distribution by Type")
        if not df_filtered_1.empty:
            def count_tests_by_rule(df_in):
                counts = {}
                for tt, grp in df_in.groupby("Type of Test"):
                    counts[tt] = int(grp["Unit No."].fillna(1).astype(
                        int).sum()) if tt == "Bioburden" else grp["Sample ID"].nunique()
                return pd.Series(counts).sort_values(ascending=False)
            test_type_dist = count_tests_by_rule(df_filtered_1)
            col1, col2 = st.columns([1, 1])
            with col1:
                import matplotlib.pyplot as plt
                total_tests_1 = int(test_type_dist.sum())
                fig, ax = plt.subplots(figsize=(10, 8))
                ax.pie(test_type_dist.values, labels=test_type_dist.index,
                       autopct=lambda pct: f'{int(round(pct/100.*total_tests_1))}\n({pct:.1f}%)' if pct > 2 else '',
                       colors=plt.cm.Set3(range(len(test_type_dist))), startangle=90, textprops={'fontsize': 9})
                ax.set_title(
                    f"Test Distribution by Type\n({start_date_1} to {end_date_1})\nTotal: {total_tests_1} tests", fontsize=12, fontweight='bold')
                st.pyplot(fig, use_container_width=True)
            with col2:
                st.metric("Total Tests", total_tests_1)
                st.caption(
                    "Bioburden = units counted individually. All other tests = unique Sample IDs.")
                for tt in test_type_dist.index:
                    count = int(test_type_dist[tt])
                    if count > 0:
                        st.write(
                            f"  • {'(units)' if tt == 'Bioburden' else '(unique IDs)'} {tt}: **{count}**")
        else:
            st.info("No data available for the selected date range")

        st.markdown("---")
        st.subheader("Chart 2: Monthly Testing Volume by Test Type")
        if not df_filtered_2.empty:
            df2 = df_filtered_2.copy()
            df2["Month"] = df2["Received Date"].dt.strftime("%b %Y")
            months_ordered = sorted(df2["Month"].unique(
            ), key=lambda m: pd.to_datetime(m, format="%b %Y"))
            test_types_present = df2["Type of Test"].unique()
            monthly_dict = {}
            for month in months_ordered:
                mdf = df2[df2["Month"] == month]
                monthly_dict[month] = {}
                for tt in test_types_present:
                    grp = mdf[mdf["Type of Test"] == tt]
                    monthly_dict[month][tt] = int(grp["Unit No."].fillna(1).astype(
                        int).sum()) if tt == "Bioburden" else grp["Sample ID"].nunique()
            monthly_test_counts = pd.DataFrame(
                monthly_dict).T.fillna(0).astype(int)
            col1, col2 = st.columns([1, 1])
            with col1:
                import matplotlib.pyplot as plt
                from matplotlib.ticker import MaxNLocator
                fig, ax = plt.subplots(figsize=(12, 6))
                x = range(len(monthly_test_counts.index))
                width = 0.15
                colors = plt.cm.Set3(range(len(monthly_test_counts.columns)))
                for i, tt in enumerate(monthly_test_counts.columns):
                    offset = (i - len(monthly_test_counts.columns) /
                              2 + 0.5) * width
                    bars = ax.bar([pos + offset for pos in x], monthly_test_counts[tt],
                                  width=width, label=tt, color=colors[i], edgecolor='black', linewidth=0.5)
                    for bar in bars:
                        h = bar.get_height()
                        if h > 0:
                            ax.text(bar.get_x() + bar.get_width()/2, h + 0.2,
                                    str(int(h)), ha='center', va='bottom', fontsize=7)
                ax.set_xticks(list(x))
                ax.set_xticklabels(monthly_test_counts.index,
                                   rotation=45, ha='right')
                ax.set_ylabel("Number of Tests")
                ax.set_xlabel("Month")
                ax.set_title(
                    f"Monthly Testing Volume by Test Type\n({start_date_2} to {end_date_2})")
                ax.legend(bbox_to_anchor=(1.05, 1),
                          loc='upper left', fontsize=9)
                ax.grid(axis='y', alpha=0.3)
                ax.yaxis.set_major_locator(MaxNLocator(integer=True))
                ax.set_ylim(
                    0, int(monthly_test_counts.values.max()) * 1.15 + 1)
                plt.tight_layout()
                st.pyplot(fig)
            with col2:
                total_unique = int(monthly_test_counts.values.sum())
                st.metric("Total Tests", total_unique)
                st.metric(
                    "Average per Month", f"{total_unique / max(len(monthly_test_counts), 1):.1f}")
                for month in monthly_test_counts.index:
                    st.write(f"**{month}:**")
                    for tt in monthly_test_counts.columns:
                        count = monthly_test_counts.loc[month, tt]
                        if count > 0:
                            st.write(f"  • {tt}: **{int(count)}**")
        else:
            st.info("No data available for the selected date range")

        st.markdown("---")
        st.subheader("📋 All Samples")
        df_display = df.copy()
        df_display["Received Date"] = df_display["Received Date"].dt.strftime(
            "%d/%m/%Y")
        display_cols = [c for c in ["Sample ID", "Unit No.", "Received Date", "Sample Type", "Sample Batch No.",
                                    "Customer Name", "Type of Test", "Test Status", "Product Name", "Reference No."] if c in df_display.columns]
        st.dataframe(df_display[display_cols], use_container_width=True)

        import io
        excel_buffer = io.BytesIO()
        df_download = df.copy()
        df_download["Received Date"] = df_download["Received Date"].dt.strftime(
            "%d/%m/%Y")
        df_download.to_excel(excel_buffer, index=False, engine="openpyxl")
        excel_buffer.seek(0)
        try:
            df_download.to_excel(os.path.join(
                BACKUP_DIR, f"Database1803_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"), index=False, engine="openpyxl")
        except Exception:
            pass
        st.download_button("📥 Download Full Database as Excel", data=excel_buffer, file_name="Database1803.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_full_database")

        st.markdown("---")
        st.subheader("🗑️ Delete Samples")
        col1, col2 = st.columns(2)
        with col1:
            del_id = st.text_input("Sample ID to delete:")
            del_unit = st.number_input(
                "Unit No. to delete (0 = all units)", min_value=0, value=0, step=1)
            if st.button("Delete Sample", key="delete_single"):
                if del_id:
                    unit_arg = int(del_unit) if del_unit != 0 else None
                    deleted = delete_rows_targeted(del_id, unit_no=unit_arg)
                    if deleted > 0:
                        st.success(f"✅ Deleted {deleted} row(s)!")
                        st.rerun()
                    else:
                        st.error("❌ Sample ID not found")
                else:
                    st.warning("Please enter a Sample ID")
        with col2:
            if st.button("Delete ALL Samples", key="delete_all"):
                if st.checkbox("⚠️ I confirm deletion of ALL samples", key="confirm_delete_all"):
                    save_data(pd.DataFrame(columns=COLUMNS))
                    st.success("✅ All samples deleted!")
                    st.rerun()

# =====================================================
# ADD SAMPLE
# =====================================================
elif menu == "Add Sample":
    st.subheader("Add Samples (Batch Entry)")
    if "batch_samples" not in st.session_state:
        st.session_state.batch_samples = []

    received_date = st.date_input("Received Date", value=date.today())
    is_existing = st.radio("Is this another unit of an existing Sample ID?", [
                           "No (New Sample)", "Yes (Existing Sample ID)"])

    df_existing = get_data()
    if not df_existing.empty:
        df_existing["Received Date"] = pd.to_datetime(
            df_existing["Received Date"], errors="coerce")

    year = received_date.year
    month = received_date.month

    if is_existing == "No (New Sample)":
        serial = get_next_serial(df_existing, st.session_state.batch_samples)
        sample_id = f"MIC-{serial:04d}-{month:02d}-{year}"
        st.text_input("Sample ID", value=sample_id, disabled=True)
    else:
        existing_ids = df_existing["Sample ID"].dropna().unique().tolist()
        sample_id = st.selectbox("Select Existing Sample ID", existing_ids)
        if sample_id:
            existing_units = df_existing[df_existing["Sample ID"] == sample_id][[
                "Unit No.", "Sample Batch No.", "Sample Type"]].copy()
            if not existing_units.empty:
                st.markdown("**Existing units for this Sample ID:**")
                st.dataframe(existing_units,
                             use_container_width=True, hide_index=True)

    test_type_option = st.selectbox("Type of Test", TEST_TYPES)
    if test_type_option == "Other (Not Listed)":
        test_type = st.text_input("Please specify Test Type")
    else:
        test_type = test_type_option

    sample_type_option = st.selectbox("Sample Type", SAMPLE_TYPES)
    if sample_type_option == "Other (Not listed)":
        sample_type = st.text_input("Please specify Sample Type")
    else:
        sample_type = sample_type_option

    col_ar, col_en = st.columns(2)
    with col_ar:
        saved_ar = get_customers_ar()
        if saved_ar:
            ar_choice = st.selectbox(
                "Customer Name (Arabic) - السادة", ["-- Type new --"] + saved_ar, key="ar_select")
            customer_name_ar = st.text_input(
                "Enter new Arabic Customer Name", key="ar_new") if ar_choice == "-- Type new --" else ar_choice
        else:
            customer_name_ar = st.text_input("Customer Name (Arabic) - السادة")

    with col_en:
        saved_en = get_customers_en()
        if saved_en:
            en_choice = st.selectbox("Customer Name (English)", [
                                     "-- Type new --"] + saved_en, key="en_select")
            customer_name_en = st.text_input(
                "Enter new English Customer Name", key="en_new") if en_choice == "-- Type new --" else en_choice
        else:
            customer_name_en = st.text_input("Customer Name (English)")

    customer_name = customer_name_en or customer_name_ar

    if test_type != "Environmental":
        sample_batch_no = st.text_input("Sample Batch No.")
        has_ref = st.radio("Does this sample have Reference No.?", [
                           "No", "Yes"], horizontal=True)
        ref_no = st.text_input("Reference No.") if has_ref == "Yes" else ""
        num_units = st.number_input(
            "Number of Units to add", min_value=1, max_value=100, value=1, step=1)
    else:
        sample_batch_no = ""
        ref_no = ""
        num_units = 1

    if test_type == "Environmental":
        st.subheader("📊 Environmental Test - Multiple Samples")
        num_samples = st.number_input(
            "How many samples do you have?", min_value=1, max_value=50, value=1, step=1)
        if "env_samples_data" not in st.session_state:
            st.session_state.env_samples_data = []
        if len(st.session_state.env_samples_data) != num_samples:
            st.session_state.env_samples_data = [
                {"Product Name": ""} for _ in range(num_samples)]

        env_table_data = []
        for i in range(num_samples):
            pn = st.text_input(
                f"Sample {i+1} - Product Name", value=st.session_state.env_samples_data[i]["Product Name"], key=f"env_product_{i}")
            st.session_state.env_samples_data[i]["Product Name"] = pn
            env_table_data.append({"Sample #": i + 1, "Product Name": pn})
        st.dataframe(pd.DataFrame(env_table_data),
                     use_container_width=True, hide_index=True)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("➕ Add Environmental Samples to Batch"):
                if all(s["Product Name"].strip() for s in st.session_state.env_samples_data):
                    base_serial = get_next_serial(
                        df_existing, st.session_state.batch_samples)
                    for i, env_sample in enumerate(st.session_state.env_samples_data):
                        st.session_state.batch_samples.append({
                            "Received Date": received_date, "Sample ID": f"MIC-{(base_serial+i):04d}-{month:02d}-{year}",
                            "Unit No.": 1, "Sample Type": sample_type, "Sample Batch No.": sample_batch_no,
                            "Customer Name": customer_name, "Customer Name (AR)": customer_name_ar,
                            "Customer Name (EN)": customer_name_en, "Reference No.": ref_no,
                            "Type of Test": test_type, "Test Performing Date": "",
                            "Test Status": "On Hold", "Product Name": env_sample["Product Name"]
                        })
                    st.session_state.env_samples_data = []
                    st.success(
                        f"✅ Added {num_samples} environmental samples to batch")
                else:
                    st.error("❌ Please fill in Product Name for all samples")
        with col2:
            if st.button("💾 Save Environmental Samples"):
                if st.session_state.batch_samples:
                    if customer_name_ar:
                        add_custom_value("CustomersAR", customer_name_ar)
                    if customer_name_en:
                        add_custom_value("CustomersEN", customer_name_en)
                    append_rows(pd.DataFrame(st.session_state.batch_samples))
                    st.session_state.batch_samples = []
                    st.session_state.env_samples_data = []
                    st.success("✅ Saved to Google Sheets")
                    st.rerun()
                else:
                    st.warning("No samples in batch to save")

        if st.session_state.batch_samples:
            st.dataframe(pd.DataFrame(st.session_state.batch_samples),
                         use_container_width=True, hide_index=True)
        st.stop()

    col1, col2 = st.columns(2)
    with col1:
        if st.button("➕ Add Sample to Batch"):
            df_reload = get_data()
            if test_type_option == "Other (Not Listed)" and test_type:
                add_custom_value("TestTypes", test_type)
            if sample_type_option == "Other (Not listed)" and sample_type:
                add_custom_value("SampleTypes", sample_type)
            if customer_name_ar:
                add_custom_value("CustomersAR", customer_name_ar)
            if customer_name_en:
                add_custom_value("CustomersEN", customer_name_en)
            if is_existing == "No (New Sample)":
                final_serial = get_next_serial(
                    df_reload, st.session_state.batch_samples)
                final_sample_id = f"MIC-{final_serial:04d}-{month:02d}-{year}"
                for unit in range(1, int(num_units) + 1):
                    st.session_state.batch_samples.append({
                        "Received Date": received_date, "Sample ID": final_sample_id, "Unit No.": unit,
                        "Sample Type": sample_type, "Sample Batch No.": sample_batch_no,
                        "Customer Name": customer_name, "Customer Name (AR)": customer_name_ar,
                        "Customer Name (EN)": customer_name_en, "Reference No.": ref_no,
                        "Type of Test": test_type, "Test Performing Date": "", "Test Status": "On Hold", "Product Name": ""
                    })
                st.success(f"✅ Added {final_sample_id} Units 1–{num_units}" if num_units >
                           1 else f"✅ Added {final_sample_id} Unit 1")
            else:
                next_unit = get_next_unit_no(
                    df_reload, sample_id, sample_batch_no)
                queued_units = [s["Unit No."] for s in st.session_state.batch_samples if s["Sample ID"]
                                == sample_id and str(s["Sample Batch No."]) == str(sample_batch_no)]
                if queued_units:
                    next_unit = max(max(queued_units) + 1, next_unit)
                for i in range(int(num_units)):
                    st.session_state.batch_samples.append({
                        "Received Date": received_date, "Sample ID": sample_id, "Unit No.": next_unit + i,
                        "Sample Type": sample_type, "Sample Batch No.": sample_batch_no,
                        "Customer Name": customer_name, "Customer Name (AR)": customer_name_ar,
                        "Customer Name (EN)": customer_name_en, "Reference No.": ref_no,
                        "Type of Test": test_type, "Test Performing Date": "", "Test Status": "On Hold", "Product Name": ""
                    })
                last_unit = next_unit + int(num_units) - 1
                st.success(f"✅ Added {sample_id} Units {next_unit}–{last_unit}" if num_units >
                           1 else f"✅ Added {sample_id} Unit {next_unit}")

    with col2:
        if st.button("💾 Save All Samples"):
            if st.session_state.batch_samples:
                append_rows(pd.DataFrame(st.session_state.batch_samples))
                st.session_state.batch_samples = []
                st.success("✅ All samples saved to Google Sheets")
                st.rerun()
            else:
                st.warning("No samples in batch")

    if st.session_state.batch_samples:
        st.subheader("Samples in current batch")
        batch_df = pd.DataFrame(st.session_state.batch_samples)
        show_cols = [c for c in ["Sample ID", "Unit No.", "Sample Type",
                                 "Sample Batch No.", "Type of Test", "Customer Name"] if c in batch_df.columns]
        st.dataframe(batch_df[show_cols],
                     use_container_width=True, hide_index=True)

# =====================================================
# EDIT SAMPLE
# =====================================================
elif menu == "Edit Sample":
    df = get_data()
    if df.empty:
        st.warning("No samples in the system. Add samples first.")
        st.stop()

    st.subheader("🔍 Search & Select Sample to Edit")
    col1, col2, col3 = st.columns(3)
    with col1:
        search_id = st.text_input(
            "Search by Sample ID", placeholder="e.g. MIC-0001-01-2025")
    with col2:
        search_customer = st.text_input("Search by Customer Name")
    with col3:
        search_status = st.selectbox("Filter by Test Status", [
                                     "All", "On Hold", "In Progress", "Released"])

    df_filtered = df.copy()
    if search_id:
        df_filtered = df_filtered[df_filtered["Sample ID"].astype(
            str).str.contains(search_id, case=False, na=False)]
    if search_customer:
        df_filtered = df_filtered[df_filtered["Customer Name"].astype(
            str).str.contains(search_customer, case=False, na=False)]
    if search_status != "All":
        df_filtered = df_filtered[df_filtered["Test Status"] == search_status]

    if df_filtered.empty:
        st.info("No samples match your search criteria.")
        st.stop()

    df_display = df_filtered.copy()
    df_display["Received Date"] = pd.to_datetime(
        df_display["Received Date"], errors="coerce").dt.strftime("%d/%m/%Y")
    st.markdown(f"**{len(df_filtered)} sample(s) found:**")
    disp_cols = [c for c in ["Sample ID", "Unit No.", "Received Date", "Customer Name", "Sample Type",
                             "Type of Test", "Test Status", "Sample Batch No.", "Reference No.", "Product Name"] if c in df_display.columns]
    st.dataframe(df_display[disp_cols],
                 use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("✏️ Edit Sample Details")
    df_filtered_copy = df_filtered.copy()
    df_filtered_copy["_label"] = df_filtered_copy.apply(
        lambda r: f"{r['Sample ID']}  |  Unit {int(r['Unit No.'])}", axis=1)
    label_to_idx = {row["_label"]: idx for idx,
                    row in df_filtered_copy.iterrows()}
    selected_label = st.selectbox("Select Sample to Edit", list(
        label_to_idx.keys()), key="edit_select_label")

    if selected_label:
        row_idx = label_to_idx[selected_label]
        row = df.loc[row_idx]
        selected_id = row["Sample ID"]
        selected_unit = int(row["Unit No."])
        row_mask = df.index == row_idx

        with st.form("edit_sample_form"):
            st.markdown(f"**Editing: `{selected_id}` — Unit {selected_unit}**")
            col1, col2 = st.columns(2)
            with col1:
                received_val = pd.to_datetime(
                    row["Received Date"], errors="coerce")
                new_received_date = st.date_input("Received Date", value=date.today(
                ) if pd.isna(received_val) else received_val.date())
                current_sample_type = str(row["Sample Type"]) if pd.notna(
                    row["Sample Type"]) else SAMPLE_TYPES[0]
                sample_type_idx = SAMPLE_TYPES.index(
                    current_sample_type) if current_sample_type in SAMPLE_TYPES else len(SAMPLE_TYPES) - 1
                new_sample_type_opt = st.selectbox(
                    "Sample Type", SAMPLE_TYPES, index=sample_type_idx)
                new_sample_type = st.text_input(
                    "Specify Sample Type", value=current_sample_type) if new_sample_type_opt == "Other (Not listed)" else new_sample_type_opt
                current_customer_en = str(row.get("Customer Name (EN)", row.get("Customer Name", ""))) if pd.notna(
                    row.get("Customer Name (EN)", row.get("Customer Name", ""))) else ""
                saved_en_edit = ["-- Type new --"] + get_customers_en()
                en_edit_choice = st.selectbox("Customer Name (English)", saved_en_edit, index=saved_en_edit.index(
                    current_customer_en) if current_customer_en in saved_en_edit else 0, key="edit_en_select")
                new_customer_en = st.text_input("Enter new English Customer Name", value=current_customer_en,
                                                key="edit_en_new") if en_edit_choice == "-- Type new --" else en_edit_choice
                new_batch = st.text_input("Sample Batch No.", value=str(
                    row["Sample Batch No."]) if pd.notna(row["Sample Batch No."]) else "")
                new_product = st.text_input("Product Name", value=str(
                    row["Product Name"]) if pd.notna(row.get("Product Name", "")) else "")
                new_unit_no = st.number_input(
                    "Unit No.", min_value=1, value=selected_unit, step=1)
            with col2:
                test_date_val = pd.to_datetime(
                    row["Test Performing Date"], errors="coerce")
                new_test_date = st.date_input("Test Performing Date", value=date.today(
                ) if pd.isna(test_date_val) else test_date_val.date())
                current_test_type = str(row["Type of Test"]) if pd.notna(
                    row["Type of Test"]) else TEST_TYPES[0]
                test_type_idx = TEST_TYPES.index(
                    current_test_type) if current_test_type in TEST_TYPES else len(TEST_TYPES) - 1
                new_test_type_opt = st.selectbox(
                    "Type of Test", TEST_TYPES, index=test_type_idx)
                new_test_type = st.text_input(
                    "Specify Test Type", value=current_test_type) if new_test_type_opt == "Other (Not Listed)" else new_test_type_opt
                current_customer_ar = str(row.get("Customer Name (AR)", "")) if pd.notna(
                    row.get("Customer Name (AR)", "")) else ""
                saved_ar_edit = ["-- Type new --"] + get_customers_ar()
                ar_edit_choice = st.selectbox("Customer Name (Arabic) - السادة", saved_ar_edit, index=saved_ar_edit.index(
                    current_customer_ar) if current_customer_ar in saved_ar_edit else 0, key="edit_ar_select")
                new_customer_ar = st.text_input("Enter new Arabic Customer Name", value=current_customer_ar,
                                                key="edit_ar_new") if ar_edit_choice == "-- Type new --" else ar_edit_choice
                new_ref = st.text_input("Reference No.", value=str(
                    row["Reference No."]) if pd.notna(row["Reference No."]) else "")
                status_options = ["On Hold", "In Progress", "Released"]
                current_status = str(row["Test Status"]) if pd.notna(
                    row["Test Status"]) else "On Hold"
                new_status = st.selectbox("Test Status", status_options, index=status_options.index(
                    current_status) if current_status in status_options else 0)

            if st.form_submit_button("💾 Save Changes", use_container_width=True):
                new_customer = new_customer_en or new_customer_ar
                if new_sample_type and new_sample_type not in SAMPLE_TYPES_DEFAULT:
                    add_custom_value("SampleTypes", new_sample_type)
                if new_test_type and new_test_type not in TEST_TYPES:
                    add_custom_value("TestTypes", new_test_type)
                if new_customer_en:
                    add_custom_value("CustomersEN", new_customer_en)
                if new_customer_ar:
                    add_custom_value("CustomersAR", new_customer_ar)
                df.loc[row_mask, "Received Date"] = pd.Timestamp(
                    new_received_date)
                df.loc[row_mask, "Unit No."] = new_unit_no
                df.loc[row_mask, "Sample Type"] = new_sample_type
                df.loc[row_mask, "Customer Name"] = new_customer
                df.loc[row_mask, "Customer Name (EN)"] = new_customer_en
                df.loc[row_mask, "Customer Name (AR)"] = new_customer_ar
                df.loc[row_mask, "Sample Batch No."] = new_batch
                df.loc[row_mask, "Reference No."] = new_ref
                df.loc[row_mask, "Type of Test"] = new_test_type
                df.loc[row_mask, "Test Performing Date"] = pd.Timestamp(
                    new_test_date).strftime("%Y-%m-%d")
                df.loc[row_mask, "Test Status"] = new_status
                df.loc[row_mask, "Product Name"] = new_product
                update_rows_targeted(df[row_mask])
                st.success(
                    f"✅ `{selected_id}` Unit {new_unit_no} updated successfully!")
                st.rerun()

        st.markdown("---")
        st.subheader("➕ Add More Units to This Batch")
        batch_no = str(row["Sample Batch No."]) if pd.notna(
            row["Sample Batch No."]) else ""
        if batch_no:
            batch_rows = df[(df["Sample ID"] == selected_id) & (
                df["Sample Batch No."].astype(str) == batch_no)].copy()
            batch_rows["Received Date"] = pd.to_datetime(
                batch_rows["Received Date"], errors="coerce").dt.strftime("%d/%m/%Y")
            st.markdown(
                f"**Current units for `{selected_id}` — Batch `{batch_no}`:**")
            b_cols = [c for c in ["Sample ID", "Unit No.", "Received Date", "Sample Type", "Sample Batch No.",
                                  "Customer Name", "Type of Test", "Test Status", "Product Name"] if c in batch_rows.columns]
            st.dataframe(batch_rows[b_cols],
                         use_container_width=True, hide_index=True)
            st.info(
                f"ℹ️ Next unit will be **Unit {get_next_unit_no(df, selected_id, batch_no)}**")
        else:
            st.info(
                "This sample has no Batch No. assigned. Please add one above first.")

        if "new_batch_samples" not in st.session_state:
            st.session_state.new_batch_samples = []

        with st.form("add_to_batch_form"):
            col1, col2 = st.columns(2)
            with col1:
                current_sample_type = str(row["Sample Type"]) if pd.notna(
                    row["Sample Type"]) else SAMPLE_TYPES[0]
                nb_st_idx = SAMPLE_TYPES.index(
                    current_sample_type) if current_sample_type in SAMPLE_TYPES else len(SAMPLE_TYPES) - 1
                nb_st_opt = st.selectbox(
                    "Sample Type", SAMPLE_TYPES, index=nb_st_idx, key="nb_sample_type")
                nb_sample_type = st.text_input(
                    "Specify Sample Type", key="nb_sample_type_other") if nb_st_opt == "Other (Not listed)" else nb_st_opt
                nb_batch_no = st.text_input(
                    "Sample Batch No.", value=batch_no, key="nb_batch_no")
                nb_product = st.text_input("Product Name", key="nb_product")
                nb_num_units = st.number_input(
                    "Number of Units to add", min_value=1, max_value=100, value=1, step=1, key="nb_num_units")
            with col2:
                nb_received_date = st.date_input(
                    "Received Date", value=date.today(), key="nb_received_date")
                current_test_type = str(row["Type of Test"]) if pd.notna(
                    row["Type of Test"]) else TEST_TYPES[0]
                nb_tt_idx = TEST_TYPES.index(
                    current_test_type) if current_test_type in TEST_TYPES else len(TEST_TYPES) - 1
                nb_tt_opt = st.selectbox(
                    "Type of Test", TEST_TYPES, index=nb_tt_idx, key="nb_test_type")
                nb_test_type = st.text_input(
                    "Specify Test Type", key="nb_test_type_other") if nb_tt_opt == "Other (Not Listed)" else nb_tt_opt
                nb_ref = st.text_input("Reference No.", key="nb_ref")

            if st.form_submit_button("➕ Add to Batch Queue", use_container_width=True):
                if nb_sample_type and nb_sample_type not in SAMPLE_TYPES_DEFAULT:
                    add_custom_value("SampleTypes", nb_sample_type)
                if nb_test_type and nb_test_type not in TEST_TYPES:
                    add_custom_value("TestTypes", nb_test_type)
                df_reload = get_data()
                next_unit = get_next_unit_no(
                    df_reload, selected_id, nb_batch_no)
                queued_units = [s["Unit No."] for s in st.session_state.new_batch_samples if s["Sample ID"]
                                == selected_id and str(s["Sample Batch No."]) == str(nb_batch_no)]
                if queued_units:
                    next_unit = max(max(queued_units) + 1, next_unit)
                for i in range(int(nb_num_units)):
                    st.session_state.new_batch_samples.append({
                        "Received Date": nb_received_date, "Sample ID": selected_id, "Unit No.": next_unit + i,
                        "Sample Type": nb_sample_type, "Sample Batch No.": nb_batch_no,
                        "Customer Name": str(row.get("Customer Name", "")),
                        "Customer Name (EN)": str(row.get("Customer Name (EN)", "")),
                        "Customer Name (AR)": str(row.get("Customer Name (AR)", "")),
                        "Reference No.": nb_ref, "Type of Test": nb_test_type,
                        "Test Performing Date": "", "Test Status": "On Hold", "Product Name": nb_product
                    })
                last_unit = next_unit + int(nb_num_units) - 1
                st.success(f"✅ Units {next_unit}–{last_unit} queued!" if nb_num_units >
                           1 else f"✅ Unit {next_unit} queued!")

        if st.session_state.new_batch_samples:
            queue_df = pd.DataFrame(st.session_state.new_batch_samples)
            q_cols = [c for c in ["Sample ID", "Unit No.", "Sample Type", "Sample Batch No.",
                                  "Type of Test", "Product Name", "Test Status"] if c in queue_df.columns]
            st.dataframe(queue_df[q_cols],
                         use_container_width=True, hide_index=True)
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Save All to Batch", use_container_width=True):
                    df_reload = get_data()
                    batch_no_val = st.session_state.new_batch_samples[0]["Sample Batch No."]
                    sid_val = st.session_state.new_batch_samples[0]["Sample ID"]
                    batch_indices = df_reload[(df_reload["Sample Batch No."].astype(str) == str(
                        batch_no_val)) & (df_reload["Sample ID"] == sid_val)].index.tolist()
                    new_rows_df = pd.DataFrame(
                        st.session_state.new_batch_samples)
                    if batch_indices:
                        last_idx = max(batch_indices)
                        df_final = pd.concat(
                            [df_reload.iloc[:last_idx+1], new_rows_df, df_reload.iloc[last_idx+1:]], ignore_index=True)
                    else:
                        df_final = pd.concat(
                            [df_reload, new_rows_df], ignore_index=True)
                    append_rows(new_rows_df)
                    st.session_state.new_batch_samples = []
                    st.success("✅ New units saved!")
                    st.rerun()
            with col2:
                if st.button("🗑️ Clear Queue", use_container_width=True):
                    st.session_state.new_batch_samples = []
                    st.rerun()

# =====================================================
# PERFORM TEST
# =====================================================
elif menu == "Perform Test":
    df = get_data()
    if df.empty:
        st.warning("No samples in the system. Add samples first.")
        st.stop()

    st.subheader("🧪 Start Test for Multiple Samples")

    selected_sample_ids = st.multiselect(
        "Select Sample ID(s) to start testing",
        df["Sample ID"].dropna().unique().tolist(),
        default=[],
    )
    test_date = st.date_input("Test Performing Date", value=date.today())

    # ── Media Suggestion & Consumption Block ───────────────────────────────
    media_consumption_inputs = {}   # {keywords_tuple: volume_ml}

    if selected_sample_ids:
        # Derive test type from first selected sample
        first_row = df[df["Sample ID"] == selected_sample_ids[0]].iloc[0]
        test_type_detected = str(first_row.get("Type of Test", "")).strip()

        st.markdown("---")
        st.markdown(f"### 🧫 Suggested Media — **{test_type_detected}** Test")

        suggestions = get_suggested_media_for_test(test_type_detected)

        if not suggestions:
            if test_type_detected == "Endotoxin":
                st.info("ℹ️ Endotoxin testing uses LAL reagents — no culture media required.")
            else:
                st.info(
                    "ℹ️ No media mapping is defined for this test type. "
                    "You can add one to `TEST_MEDIA_MAP` in app.py."
                )
        else:
            for sugg in suggestions:
                col_info, col_vol = st.columns([3, 1])
                with col_info:
                    if sugg["has_stock"]:
                        batch = sugg["latest_batch"]
                        lot       = batch["Lot No."]
                        remaining = sugg["remaining_ml"]
                        expiry    = batch["Expiry Date"]
                        expiry_str   = expiry.strftime("%Y-%m-%d") if pd.notna(expiry) else "N/A"
                        prep_date_str = batch["Date"].strftime("%Y-%m-%d") if pd.notna(batch["Date"]) else "N/A"
                        st.success(
                            f"✅ **{sugg['label']}**  \n"
                            f"📦 Media: **{batch['Media Type']}** | Lot/Batch No.: `{lot}`  \n"
                            f"🗓️ Prepared: {prep_date_str} | Expires: {expiry_str}  \n"
                            f"💧 Remaining in this batch: **{remaining:.1f} mL**"
                        )
                    else:
                        st.error(
                            f"❌ **{sugg['label']}**  \n"
                            f"⚠️ **No prepared batch available with remaining volume!**  \n"
                            f"👉 Go to the **Store App → Prepare Media** to prepare more before starting this test."
                        )
                with col_vol:
                    if sugg["has_stock"]:
                        default_ml = DEFAULT_MEDIA_CONSUMPTION_ML.get(test_type_detected, 15)
                        vol = st.number_input(
                            "Volume to use (mL)",
                            min_value=0.0,
                            max_value=float(sugg["remaining_ml"]),
                            value=min(float(default_ml), float(sugg["remaining_ml"])),
                            step=5.0,
                            key=f"media_vol_{'_'.join(str(k) for k in sugg['keywords'][:2])}",
                        )
                        media_consumption_inputs[tuple(sugg["keywords"])] = vol

        st.markdown("---")

    # ── Action Button ──────────────────────────────────────────────────────
    if st.button("▶️ Mark Test as Started", type="primary"):
        if not selected_sample_ids:
            st.warning("Please select at least one Sample ID.")
        else:
            # 1. Check that all required media have stock before proceeding
            all_media_ok = True
            for sugg in (suggestions if selected_sample_ids else []):
                if not sugg["has_stock"]:
                    st.error(
                        f"❌ Cannot start test — **{sugg['label']}** has no prepared batch available. "
                        "Please prepare the required media first."
                    )
                    all_media_ok = False

            if all_media_ok:
                # 2. Update test status in Google Sheet
                for sid in selected_sample_ids:
                    df.loc[df["Sample ID"] == sid, ["Test Performing Date", "Test Status"]] = [
                        test_date.strftime("%Y-%m-%d"), "In Progress"
                    ]
                update_rows_targeted(df[df["Sample ID"].isin(selected_sample_ids)])

                # 3. Deduct media consumption from media_preparation.xlsx
                deduction_messages = []
                for keywords_tuple, volume_ml in media_consumption_inputs.items():
                    if volume_ml > 0:
                        ok, msg = deduct_media_consumption(list(keywords_tuple), volume_ml)
                        deduction_messages.append((ok, msg, volume_ml, keywords_tuple))

                st.success(f"✅ {len(selected_sample_ids)} test(s) marked as **In Progress**")

                if deduction_messages:
                    st.markdown("**🧫 Media Consumption Summary:**")
                    for ok, msg, vol_ml, kw in deduction_messages:
                        if ok:
                            st.info(f"• Deducted **{vol_ml:.1f} mL** — {msg}")
                        else:
                            st.warning(f"• ⚠️ Could not deduct media: {msg}")

                # 4. Re-check remaining after deduction and warn if now empty
                st.markdown("**📊 Updated Media Stock:**")
                if selected_sample_ids:
                    refreshed = get_suggested_media_for_test(test_type_detected)
                    for sugg in refreshed:
                        if sugg["has_stock"]:
                            rem = sugg["remaining_ml"]
                            if rem <= 0:
                                st.error(
                                    f"🚨 **{sugg['label']}** batch is now fully consumed! "
                                    "Please prepare more before the next test."
                                )
                            elif rem < 50:
                                st.warning(
                                    f"⚠️ **{sugg['label']}** batch is running low: only **{rem:.1f} mL** remaining."
                                )
                            else:
                                st.success(f"✅ **{sugg['label']}**: {rem:.1f} mL remaining.")
                        else:
                            st.error(
                                f"🚨 **{sugg['label']}** — No batch remaining! "
                                "Go to **Store App → Prepare Media** to prepare more."
                            )

# =====================================================
# ENTER RESULTS
# =====================================================
elif menu == "Enter Results":
    df = get_data()
    if df.empty:
        st.warning("No samples in the system. Add samples first.")
        st.stop()

    sample_id_selected = st.selectbox(
        "Select Sample ID", df["Sample ID"].dropna().unique().tolist())
    sample_rows = df[df["Sample ID"] == sample_id_selected]
    sample_row = sample_rows.iloc[0]
    test_type = sample_row["Type of Test"].strip()
    report_prefix = safe_report_filename(sample_id_selected)

    if sample_row["Test Status"] == "On Hold":
        st.warning(
            "⚠️ This sample is On Hold. Please mark it as In Progress first (Perform Test page).")
        st.stop()

    def safe(val): return "" if pd.isna(val) else str(val)

    sample_types_comma = ", ".join(
        sample_rows["Sample Type"].dropna().astype(str).unique())
    sample_types_newline = "\n".join(
        sample_rows["Sample Type"].dropna().astype(str).unique())
    batch_list = sample_rows["Sample Batch No."].dropna().astype(str).unique()
    batch_numbers_newline = "\n".join(batch_list)
    reference_text = "\n".join(
        sample_rows["Reference No."].dropna().astype(str).unique())

    # BIOBURDEN
    if test_type == "Bioburden":
        st.subheader("Enter Test Results - Bioburden")
        issuing_date = st.date_input("Issuing Date", value=date.today())
        st.markdown("### Total Aerobic Microbial Count (CFU/ml)")
        tamc_growth = st.radio("Is there microbial growth?", [
                               "No", "Yes"], key="tamc_growth", horizontal=True)
        if tamc_growth == "Yes":
            tamc_val = st.number_input(
                "Enter TAMC value (CFU/ml)", min_value=1, step=1)
            tamc_base_text = f"{tamc_val} CFU/ml"
        else:
            tamc_base_text = "No microbial growth was detected"
        st.markdown("### Total Combined Yeasts / Molds Count (CFU/ml)")
        tymc_growth = st.radio("Is there microbial growth?", [
                               "No", "Yes"], key="tymc_growth", horizontal=True)
        if tymc_growth == "Yes":
            tymc_val = st.number_input(
                "Enter TYMC value (CFU/ml)", min_value=1, step=1)
            tymc_base_text = f"{tymc_val} CFU/ml"
        else:
            tymc_base_text = "No microbial growth was detected"

        if st.button("Generate Bioburden Report"):
            try:
                doc = Document(TEMPLATE_BIOBURDEN)
            except FileNotFoundError:
                st.error(f"Template '{TEMPLATE_BIOBURDEN}' not found.")
                st.stop()
            header_data = {
                "{{received_date}}": format_report_date(sample_row["Received Date"]),
                "{{test_performing_date}}": format_report_date(sample_row["Test Performing Date"]),
                "{{issuing_date}}": format_report_date(issuing_date),
                "{{customer_name}}": safe(sample_row["Customer Name"]),
                "{{sample_id}}": sample_id_selected, "{{sample_type}}": sample_types_comma,
                "{{sample_batch_no}}": batch_numbers_newline, "{{reference_no}}": reference_text
            }
            for p in doc.paragraphs:
                for k, v in header_data.items():
                    if k in p.text:
                        p.text = p.text.replace(k, v)
            replace_placeholders_in_tables(doc, header_data)
            sample_info_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "<<SAMPLE_INFO_TABLE>>" in cell.text:
                            sample_info_table = table
                            cell.text = sample_types_newline
                            break
                    if sample_info_table:
                        break
                if sample_info_table:
                    break
            if not sample_info_table:
                st.error("Template error: <<SAMPLE_INFO_TABLE>> not found")
                st.stop()
            results_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "<<RESULT_TABLE>>" in cell.text:
                            results_table = table
                            break
                    if results_table:
                        break
                if results_table:
                    break
            if not results_table:
                st.error("Template error: <<RESULT_TABLE>> not found")
                st.stop()
            while len(results_table.rows) > 1:
                results_table._tbl.remove(results_table.rows[1]._tr)
            for batch in batch_list:
                r = results_table.add_row().cells
                r[0].text = sample_id_selected
                r[1].text = f"{tamc_base_text} for batch number {batch}"
                r[2].text = f"{tymc_base_text} for batch number {batch}"
            doc.save(report_path(f"{report_prefix}_Bioburden_Report.docx"))
            df.loc[df["Sample ID"] == sample_id_selected,
                   "Test Status"] = "Released"
            update_rows_targeted(df[df["Sample ID"] == sample_id_selected])
            st.success("✅ Bioburden report generated successfully")

    bioburden_report_path = find_report_path(
        f"{report_prefix}_Bioburden_Report.docx")
    if bioburden_report_path:
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            with open(bioburden_report_path, "rb") as fh:
                st.download_button("📥 Download as Word (.docx)", data=fh.read(), file_name=os.path.basename(bioburden_report_path),
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_bioburden_docx")
        with col_dl2:
            pdf = convert_docx_to_pdf(bioburden_report_path)
            if pdf:
                with open(pdf, "rb") as fh:
                    st.download_button("📥 Download as PDF", data=fh.read(), file_name=os.path.basename(
                        pdf), mime="application/pdf", key="dl_bioburden_pdf")
            else:
                st.info("PDF conversion requires LibreOffice.")

    # STERILITY
    elif test_type == "Sterility":
        st.subheader("Enter Test Results - Sterility")
        issuing_date = st.date_input(
            "Issuing Date", value=date.today(), key="sterility_issue_date")
        growth = st.radio("Was there any bacterial or fungal growth?", [
                          "No", "Yes"], horizontal=True)
        user_text = st.text_area(
            "Enter sterility test result", height=150) if growth == "Yes" else ""

        if st.button("Generate Sterility Report"):
            try:
                doc = Document(TEMPLATE_STERILITY)
            except FileNotFoundError:
                st.error(f"Template '{TEMPLATE_STERILITY}' not found.")
                st.stop()
            header_data = {
                "{{customer_name}}": safe(sample_row["Customer Name"]),
                "{{received_date}}": format_report_date(sample_row["Received Date"]),
                "{{test_performing_date}}": format_report_date(sample_row["Test Performing Date"]),
                "{{issuing_date}}": format_report_date(issuing_date),
                "{{sample_type}}": sample_types_comma, "{{sample_id}}": sample_id_selected,
                "{{sample_batch_no}}": batch_numbers_newline, "{{reference_no}}": reference_text
            }
            for p in doc.paragraphs:
                for k, v in header_data.items():
                    if k in p.text:
                        p.text = p.text.replace(k, v)
            replace_placeholders_in_tables(doc, header_data)
            sample_info_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "<<SAMPLE_INFO_TABLE>>" in cell.text:
                            sample_info_table = table
                            cell.text = sample_types_newline
                            break
                    if sample_info_table:
                        break
                if sample_info_table:
                    break
            if not sample_info_table:
                st.error("Template error: <<SAMPLE_INFO_TABLE>> not found")
                st.stop()
            result_text = user_text if growth == "Yes" else "\n\n".join(
                f"{i}- No evidence of microbial growth (Bacteria and Fungi) is found, for batch number {b}. It complies with the test for sterility."
                for i, b in enumerate(batch_list, start=1))
            for p in doc.paragraphs:
                if "<<RESULT_TEXT>>" in p.text:
                    p.text = result_text
            doc.save(report_path(f"{report_prefix}_Sterility_Report.docx"))
            df.loc[df["Sample ID"] == sample_id_selected,
                   "Test Status"] = "Released"
            update_rows_targeted(df[df["Sample ID"] == sample_id_selected])
            st.success("✅ Sterility report generated successfully")

    sterility_report_path = find_report_path(
        f"{report_prefix}_Sterility_Report.docx")
    if sterility_report_path:
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            with open(sterility_report_path, "rb") as fh:
                st.download_button("📥 Download as Word (.docx)", data=fh.read(), file_name=os.path.basename(sterility_report_path),
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_sterility_docx")
        with col_dl2:
            pdf = convert_docx_to_pdf(sterility_report_path)
            if pdf:
                with open(pdf, "rb") as fh:
                    st.download_button("📥 Download as PDF", data=fh.read(), file_name=os.path.basename(
                        pdf), mime="application/pdf", key="dl_sterility_pdf")
            else:
                st.info("PDF conversion requires LibreOffice.")

    # ENDOTOXIN
    elif test_type == "Endotoxin":
        st.subheader("Enter Test Results - Endotoxin")
        issuing_date = st.date_input(
            "Issuing Date", value=date.today(), key="endotoxin_issue_date")
        endotoxin_result = st.text_input(
            "Endotoxin result (EU/ml)", placeholder="e.g. <0.01 or 0.05").strip()

        if st.button("Generate Endotoxin Report"):
            try:
                doc = Document(TEMPLATE_ENDOTOXIN)
            except FileNotFoundError:
                st.error(f"Template '{TEMPLATE_ENDOTOXIN}' not found.")
                st.stop()
            cnv = safe(sample_row["Customer Name"])
            header_data = {
                "{{customer_name}}": cnv, "{{{customer_name}": cnv,
                "{{received_date}}": format_report_date(sample_row["Received Date"]),
                "{{test_performing_date}}": format_report_date(sample_row["Test Performing Date"]),
                "{{issuing_date}}": format_report_date(issuing_date),
                "{{sample_type}}": sample_types_comma, "{{sample_id}}": sample_id_selected
            }
            for p in doc.paragraphs:
                for k, v in header_data.items():
                    if k in p.text:
                        p.text = p.text.replace(k, v)
            replace_placeholders_in_tables(doc, header_data)
            sample_info_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "<<SAMPLE_INFO_TABLE>>" in cell.text:
                            sample_info_table = table
                            cell.text = sample_types_newline
                            break
                    if sample_info_table:
                        break
                if sample_info_table:
                    break
            if not sample_info_table:
                st.error("Template error: <<SAMPLE_INFO_TABLE>> not found")
                st.stop()
            results_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "<<RESULT_TABLE>>" in cell.text:
                            results_table = table
                            break
                    if results_table:
                        break
                if results_table:
                    break
            if not results_table:
                st.error("Template error: <<RESULT_TABLE>> not found")
                st.stop()
            marker_row_idx = next((i for i, row in enumerate(results_table.rows) if any(
                "<<RESULT_TABLE>>" in c.text for c in row.cells)), None)
            if marker_row_idx is None:
                st.error("<<RESULT_TABLE>> marker not found")
                st.stop()
            for i in range(len(results_table.rows) - 1, marker_row_idx, -1):
                results_table._tbl.remove(results_table.rows[i]._tr)
            hc = results_table.rows[max(0, marker_row_idx - 1)].cells
            product_cols = [i for i, c in enumerate(
                hc) if "Product Name" in c.text]
            sample_id_cols = [i for i, c in enumerate(
                hc) if "Sample ID" in c.text]
            batch_cols = [i for i, c in enumerate(
                hc) if "Batch number" in c.text]
            endotoxin_cols = [i for i, c in enumerate(
                hc) if "Endotoxin level" in c.text]
            final_endotoxin = endotoxin_result or "Not specified"
            for idx, (_, r) in enumerate(sample_rows.iterrows()):
                cells = results_table.rows[marker_row_idx].cells if idx == 0 else results_table.add_row(
                ).cells
                if "<<RESULT_TABLE>>" in cells[0].text:
                    cells[0].text = ""
                batch_value = safe(r["Sample Batch No."])
                ref_value = safe(r["Reference No."])
                if ref_value:
                    batch_value = f"{batch_value}\n{ref_value}" if batch_value else ref_value
                for ci in product_cols or [0]:
                    if ci < len(cells):
                        cells[ci].text = safe(r["Sample Type"])
                for ci in sample_id_cols or [1]:
                    if ci < len(cells):
                        cells[ci].text = sample_id_selected
                for ci in batch_cols or [2]:
                    if ci < len(cells):
                        cell = cells[ci]
                        cell.text = ""
                        if "\n" in batch_value:
                            first, rest = batch_value.split("\n", 1)
                            cell.paragraphs[0].text = first
                            cell.add_paragraph(rest)
                        else:
                            cell.paragraphs[0].text = batch_value
                for ci in endotoxin_cols or [3]:
                    if ci < len(cells):
                        cells[ci].text = f"{final_endotoxin} "
            doc.save(report_path(f"{report_prefix}_Endotoxin_Report.docx"))
            df.loc[df["Sample ID"] == sample_id_selected,
                   "Test Status"] = "Released"
            update_rows_targeted(df[df["Sample ID"] == sample_id_selected])
            st.success("✅ Endotoxin report generated successfully")

    endotoxin_report_path = find_report_path(
        f"{report_prefix}_Endotoxin_Report.docx")
    if endotoxin_report_path:
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            with open(endotoxin_report_path, "rb") as fh:
                st.download_button("📥 Download as Word (.docx)", data=fh.read(), file_name=os.path.basename(endotoxin_report_path),
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_endotoxin_docx")
        with col_dl2:
            pdf = convert_docx_to_pdf(endotoxin_report_path)
            if pdf:
                with open(pdf, "rb") as fh:
                    st.download_button("📥 Download as PDF", data=fh.read(), file_name=os.path.basename(
                        pdf), mime="application/pdf", key="dl_endotoxin_pdf")
            else:
                st.info("PDF conversion requires LibreOffice.")

    # ENVIRONMENTAL
    elif test_type == "Environmental":
        st.subheader("Enter Test Results - Environmental")
        issuing_date = st.date_input(
            "Issuing Date", value=date.today(), key="env_issue_date")
        env_samples = df[(df["Type of Test"] == "Environmental")
                         & (df["Test Status"] == "In Progress")]
        env_sample_ids = sorted(
            env_samples["Sample ID"].dropna().unique().tolist())
        if not env_sample_ids:
            st.warning("⚠️ No Environmental samples marked as 'In Progress'.")
            st.stop()
        col1, col2 = st.columns(2)
        with col1:
            start_sample_id = st.selectbox(
                "Select Start Sample ID", env_sample_ids, key="env_start_id")
        with col2:
            end_sample_id = st.selectbox("Select End Sample ID", [
                                         s for s in env_sample_ids if s >= start_sample_id], key="env_end_id")
        sample_ids_range = generate_sample_id_range(
            start_sample_id, end_sample_id)
        if sample_ids_range:
            st.info(f"✅ Generated {len(sample_ids_range)} samples")
            results_data = []
            for idx, sid in enumerate(sample_ids_range, 1):
                col_no, col_fungi, col_bacteria = st.columns(3)
                with col_no:
                    st.write(f"#{idx}")
                with col_fungi:
                    fungi_count = st.text_input(
                        f"Fungi ({idx})", value="", key=f"fungi_{idx}")
                with col_bacteria:
                    bacteria_count = st.text_input(
                        f"Bacteria ({idx})", value="", key=f"bacteria_{idx}")
                results_data.append({"NO": idx, "Customer ID": sid, "Fungus Count (CFU)": fungi_count,
                                    "Total Bacterial Count (CFU)": bacteria_count})
            st.dataframe(pd.DataFrame(results_data),
                         use_container_width=True, hide_index=True)

            if st.button("Generate Environmental Report"):
                try:
                    doc = Document(TEMPLATE_ENVIRONMENT)
                except FileNotFoundError:
                    st.error(f"Template '{TEMPLATE_ENVIRONMENT}' not found.")
                    st.stop()
                header_data = {
                    "{{received_date}}": format_report_date(sample_row["Received Date"]),
                    "{{testing_date}}": format_report_date(sample_row["Test Performing Date"]),
                    "{{issuing_date}}": format_report_date(issuing_date),
                    "{{customer_name_ar}}": safe(sample_row.get("Customer Name (AR)", "")),
                    "{{customer_name_en}}": safe(sample_row.get("Customer Name (EN)", "")),
                    "{{sample_type}}": safe(sample_row["Sample Type"])
                }
                for p in doc.paragraphs:
                    for k, v in header_data.items():
                        if k in p.text:
                            p.text = p.text.replace(k, v)
                replace_placeholders_in_tables(doc, header_data)
                product_names_map = {sid: safe(df[df["Sample ID"] == sid].iloc[0].get(
                    "Product Name", "")) for sid in sample_ids_range if not df[df["Sample ID"] == sid].empty}
                product_table = next(
                    (t for t in doc.tables for row in t.rows for cell in row.cells if "Product Name" in cell.text), None)
                if product_table:
                    while len(product_table.rows) > 1:
                        product_table._tbl.remove(product_table.rows[1]._tr)
                    for sid in sample_ids_range:
                        r = product_table.add_row().cells
                        r[0].text = product_names_map.get(sid, "")
                        r[1].text = sid
                results_table = next(
                    (t for t in doc.tables for row in t.rows for cell in row.cells if "Total Bacterial Count" in cell.text or "Fungus Count" in cell.text), None)
                if not results_table:
                    st.error("Results table not found in template")
                    st.stop()
                while len(results_table.rows) > 1:
                    results_table._tbl.remove(results_table.rows[1]._tr)
                for result in results_data:
                    r = results_table.add_row().cells
                    r[0].text = str(result["NO"])
                    r[1].text = result["Customer ID"]
                    r[2].text = str(result["Fungus Count (CFU)"])
                    r[3].text = str(result["Total Bacterial Count (CFU)"])
                doc.save(report_path(
                    f"{report_prefix}_Environmental_Report.docx"))
                df.loc[df["Sample ID"] == sample_id_selected,
                       "Test Status"] = "Released"
                update_rows_targeted(df[df["Sample ID"] == sample_id_selected])
                st.success("✅ Environmental report generated successfully")

        env_report_path = find_report_path(
            f"{report_prefix}_Environmental_Report.docx")
        if env_report_path:
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                with open(env_report_path, "rb") as fh:
                    st.download_button("📥 Download as Word (.docx)", data=fh.read(), file_name=os.path.basename(env_report_path),
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_environment_docx")
            with col_dl2:
                pdf = convert_docx_to_pdf(env_report_path)
                if pdf:
                    with open(pdf, "rb") as fh:
                        st.download_button("📥 Download as PDF", data=fh.read(), file_name=os.path.basename(
                            pdf), mime="application/pdf", key="dl_environment_pdf")
                else:
                    st.info("PDF conversion requires LibreOffice.")
