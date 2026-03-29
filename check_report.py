from docx import Document
import os

reports_dir = "reports"
if os.path.exists(reports_dir):
    files = sorted([f for f in os.listdir(reports_dir) if f.endswith('.docx')])
    if files:
        test_file = os.path.join(reports_dir, files[0])
        doc = Document(test_file)
        
        print(f"Testing: {files[0]}")
        print(f"Total tables: {len(doc.tables)}\n")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"First table: {len(table.rows)} rows x {len(table.rows[0].cells)} cols")
            
            first_cell = table.rows[0].cells[0]
            print(f"First cell text: '{first_cell.text}'")
            
            bold_count = 0
            total_runs = 0
            for para in first_cell.paragraphs:
                for run in para.runs:
                    total_runs += 1
                    if run.bold:
                        bold_count += 1
                    print(f"  Run: '{run.text}' | Bold={run.bold}")
            
            print(f"\nBold runs: {bold_count}/{total_runs}")
    else:
        print("No .docx files in reports folder")
else:
    print("Reports folder doesn't exist")
